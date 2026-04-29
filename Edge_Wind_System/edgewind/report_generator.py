"""
报表生成模块
处理Word文档导出等功能
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime, timedelta
import logging
import io
import base64

logger = logging.getLogger(__name__)


def set_run_font(run, font_name='宋体', size=None, bold=False, color=None):
    """设置文本运行（run）的字体属性，支持中文字体"""
    run.font.name = 'Times New Roman'  # 默认西文字体
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)  # 中文字体
    if size:
        run.font.size = size
    if bold:
        run.font.bold = True
    if color:
        run.font.color.rgb = color


def set_cell_shading(cell, fill_hex):
    """
    设置表格单元格的背景颜色
    fill_hex: 颜色十六进制字符串（例如："E6F3FF"）
    """
    tc_pr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, size=Pt(11)):
    """设置单元格文本，支持中文字体（宋体）"""
    cell.paragraphs[0].clear()
    run = cell.paragraphs[0].add_run(str(text))
    set_run_font(run, size=size, bold=bold)


def calculate_health_score(devices):
    """
    计算系统健康分数（使用加权扣分模型）
    
    基础分数：100
    扣分规则：
    - 离线设备：-10分
    - 故障设备（根据故障代码）：
        - 严重故障（E04, E05）：-25分
        - 主要故障（E02, E03）：-15分
        - 警告故障（E01）：-5分
        - 未知故障：-10分
    
    Args:
        devices: 设备列表
        
    Returns:
        int: 健康分数（0-100）
    """
    score = 100
    
    for device in devices:
        if device.status == 'offline':
            score -= 10
        elif device.status == 'faulty':
            fault_code = getattr(device, 'fault_code', 'E00') or 'E00'
            
            if fault_code in ['E04', 'E05']:
                score -= 25
            elif fault_code in ['E02', 'E03']:
                score -= 15
            elif fault_code == 'E01':
                score -= 5
            elif fault_code != 'E00':
                score -= 10
    
    return max(0, score)


def get_severity_text(fault_type):
    """根据故障类型获取严重程度文本"""
    if not fault_type:
        return '一般'
    
    if '交流窜入' in fault_type or '绝缘故障' in fault_type or \
       '接地故障' in fault_type or '接地' in fault_type or \
       'IGBT' in fault_type or '开路' in fault_type:
        return '严重'
    elif '电容老化' in fault_type or '电容' in fault_type:
        return '主要'
    elif '过压' in fault_type or '过流' in fault_type:
        return '主要'
    else:
        return '一般'


def generate_report_docx(work_order, device):
    """
    生成故障诊断报告（Word文档）
    
    Args:
        work_order: 工单对象
        device: 设备对象
        
    Returns:
        Document: python-docx文档对象
    """
    try:
        doc = Document()
        
        # 标题
        title = doc.add_heading('风电场DC系统故障诊断报告', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 基本信息表格
        doc.add_heading('1. 基本信息', level=1)
        table = doc.add_table(rows=5, cols=2)
        table.style = 'Light Grid Accent 1'
        
        # 填充基本信息
        info_data = [
            ('设备编号', device.device_id),
            ('安装位置', device.location),
            ('故障时间', (work_order.fault_time + timedelta(hours=8)).strftime('%Y-%m-%d %H:%M:%S') if work_order.fault_time else '未知'),
            ('故障类型', work_order.fault_type or '未知'),
            ('严重程度', get_severity_text(work_order.fault_type))
        ]
        
        for i, (key, value) in enumerate(info_data):
            set_cell_text(table.rows[i].cells[0], key, bold=True)
            set_cell_text(table.rows[i].cells[1], value)
            set_cell_shading(table.rows[i].cells[0], 'E6F3FF')
        
        # AI诊断建议
        doc.add_heading('2. AI智能诊断', level=1)
        if work_order.ai_recommendation:
            p = doc.add_paragraph()
            run = p.add_run(work_order.ai_recommendation)
            set_run_font(run, size=Pt(10))
        else:
            doc.add_paragraph('暂无AI诊断建议')
        
        # 页脚
        doc.add_paragraph()
        footer = doc.add_paragraph('报告生成时间: ' + datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_run_font(footer.runs[0], size=Pt(9))
        
        return doc
    
    except Exception as e:
        logger.error(f"生成报告失败: {e}")
        raise


def _strip_english_in_brackets(text: str) -> str:
    """
    去除括号内英文（用于 Word 导出，避免显示过长/不符合中文要求）。
    示例：
    - “绝缘故障 (Insulation Fault)” -> “绝缘故障”
    - “交流窜入（AC Intrusion）” -> “交流窜入”
    """
    if not text:
        return ''
    s = str(text).strip()
    for sep in [' (', '（']:
        if sep in s:
            s = s.split(sep, 1)[0].strip()
            break
    return s


def _infer_fault_code_from_type(fault_type: str) -> str:
    """从故障类型文本推断故障代码（E01-E05），用于严重程度判断。"""
    ft = fault_type or ''
    if '交流窜入' in ft:
        return 'E01'
    if '绝缘' in ft:
        return 'E02'
    if '电容' in ft:
        return 'E03'
    if 'IGBT' in ft or '开路' in ft:
        return 'E04'
    if '接地' in ft:
        return 'E05'
    return 'E00'


def generate_workorder_docx(work_order, device, graph_image_dataurl: str | None = None):
    """
    生成“维修派工单”Word文档（参考旧版 app.py 的专业排版）。

    说明：
    - 文档内容尽量使用中文（符合项目规范），自动去除括号内英文。
    - 若提供 graph_image_dataurl（ECharts dataURL），会插入到附录中并居中显示。
    """
    doc = Document()

    # ==================== 页面/页边距 ====================
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    # ==================== 页眉品牌 ====================
    header = section.header
    p_brand = header.paragraphs[0]
    p_brand.text = ''
    run_brand = p_brand.add_run('EdgeWind 智能边缘监测系统｜数字化运维存档')
    set_run_font(run_brand, size=Pt(9), color=RGBColor(128, 128, 128))
    p_brand.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # ==================== 主标题 ====================
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run('设备维修派工单')
    set_run_font(run_title, size=Pt(22), bold=True, color=RGBColor(0, 51, 102))
    p_title.paragraph_format.space_after = Pt(24)

    # ==================== 工单编号（右对齐）====================
    order_id = getattr(work_order, 'id', None) or 0
    fault_time = getattr(work_order, 'fault_time', None)
    order_date = (fault_time or datetime.now()).strftime('%Y%m%d')
    order_id_str = f"WO-{order_date}-{int(order_id):03d}"
    p_id = doc.add_paragraph()
    p_id.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_id_label = p_id.add_run('工单编号：')
    set_run_font(run_id_label, size=Pt(10.5), bold=True)
    run_id_val = p_id.add_run(order_id_str)
    set_run_font(run_id_val, size=Pt(10.5), color=RGBColor(192, 0, 0))

    # ==================== 核心数据表格 ====================
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    status_map = {
        'pending': '待处理',
        'processing': '处理中',
        'resolved': '已解决',
        'fixed': '已解决'
    }

    device_id = getattr(device, 'device_id', None) or getattr(work_order, 'device_id', '') or '未知设备'
    device_location = getattr(device, 'location', None)
    order_location = getattr(work_order, 'location', None)

    device_loc_norm = str(device_location).strip() if device_location is not None else ''
    order_loc_norm = str(order_location).strip() if order_location is not None else ''
    device_id_norm = str(device_id).strip()

    # 优先使用有效的设备安装位置；若设备位置为空/等于设备名称，则回退工单位置
    if device_loc_norm and device_loc_norm != device_id_norm:
        location = device_loc_norm
    elif order_loc_norm and order_loc_norm != device_id_norm:
        location = order_loc_norm
    elif device_loc_norm or order_loc_norm:
        location = device_loc_norm or order_loc_norm
    else:
        location = '未设置'
    fault_type_raw = getattr(work_order, 'fault_type', None) or '未知故障'
    fault_type_display = _strip_english_in_brackets(fault_type_raw)
    fault_code = _infer_fault_code_from_type(fault_type_display)

    # AI 建议简单排版（保持可读）
    ai_desc = getattr(work_order, 'ai_recommendation', None) or '暂无建议'
    formatted_ai_desc = str(ai_desc).replace('建议措施:', '\n建议措施:\n').replace(' 2.', '\n2.').replace(' 1.', '\n1.')

    local_fault_time = (fault_time + timedelta(hours=8)).strftime('%Y-%m-%d %H:%M:%S') if fault_time else '未知'

    data_rows = [
        ('设备名称', device_id),
        ('安装位置', location),
        ('故障类型', fault_type_display),
        ('发生时间', local_fault_time),
        ('当前状态', status_map.get(getattr(work_order, 'status', 'pending'), '待处理')),
        ('严重程度', '严重' if fault_code in ['E04', 'E05'] else ('主要' if fault_code in ['E02', 'E03'] else ('预警' if fault_code == 'E01' else '一般'))),
        ('AI 诊断建议', formatted_ai_desc)
    ]

    for label, value in data_rows:
        row = table.add_row()
        row.height = Cm(1.0) if label != 'AI 诊断建议' else Cm(2.8)

        cell_left = row.cells[0]
        cell_left.width = Cm(3.5)
        set_cell_shading(cell_left, 'E6F3FF')
        cell_left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p_l = cell_left.paragraphs[0]
        p_l.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_l = p_l.add_run(label)
        set_run_font(run_l, size=Pt(11), bold=True)

        cell_right = row.cells[1]
        cell_right.width = Cm(13.5)
        cell_right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP if label == 'AI 诊断建议' else WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p_r = cell_right.paragraphs[0]
        p_r.alignment = WD_ALIGN_PARAGRAPH.LEFT

        if '\n' in str(value):
            for line in str(value).split('\n'):
                if line.strip():
                    run_r = p_r.add_run(line.strip() + '\n')
                    set_run_font(run_r, size=Pt(10.5))
        else:
            run_r = p_r.add_run(str(value))
            set_run_font(run_r, size=Pt(10.5))

    doc.add_paragraph().paragraph_format.space_before = Pt(12)

    # ==================== 附录：故障诊断知识图谱（从第二页开始）====================
    if isinstance(graph_image_dataurl, str) and graph_image_dataurl.startswith('data:image'):
        try:
            b64 = graph_image_dataurl.split(',', 1)[1] if ',' in graph_image_dataurl else None
            if b64:
                image_data = base64.b64decode(b64)
                image_stream = io.BytesIO(image_data)

                # 关键：附录固定从第二页开始，避免与正文挤在同一页导致版式不稳定
                doc.add_page_break()

                p_graph_title = doc.add_paragraph()
                run_graph_title = p_graph_title.add_run('附录：故障诊断知识图谱')
                set_run_font(run_graph_title, size=Pt(12), bold=True)

                p_img = doc.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_img = p_img.add_run()
                # 关键：限制宽度，避免图片撑破页面（旧版是 15cm）
                run_img.add_picture(image_stream, width=Cm(15))
        except Exception as img_error:
            logger.warning(f"插入知识图谱图片失败（已忽略）：{img_error}")

    doc.add_paragraph().paragraph_format.space_before = Pt(12)

    # ==================== 签字栏（无边框表格对齐）====================
    sig_table = doc.add_table(rows=2, cols=2)
    sig_table.autofit = True
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    row1 = sig_table.rows[0]
    p1 = row1.cells[0].paragraphs[0]
    r1 = p1.add_run('维修人员签字：')
    set_run_font(r1, size=Pt(11), bold=True)

    p2 = row1.cells[1].paragraphs[0]
    r2 = p2.add_run('验收主管签字：')
    set_run_font(r2, size=Pt(11), bold=True)

    row2 = sig_table.rows[1]
    row2.height = Cm(1.5)
    row2.cells[0].paragraphs[0].add_run('__________________________')
    row2.cells[1].paragraphs[0].add_run('__________________________')

    # ==================== 页脚时间 ====================
    p_footer = doc.add_paragraph()
    p_footer.paragraph_format.space_before = Pt(20)
    p_footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    now_str = datetime.now().strftime('%Y-%m-%d %H:%M')
    run_foot = p_footer.add_run(f'打印时间：{now_str}｜由 EdgeWind 系统生成')
    set_run_font(run_foot, size=Pt(9), color=RGBColor(100, 100, 100))

    return doc
